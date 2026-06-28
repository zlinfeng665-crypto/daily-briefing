#!/usr/bin/env python3
"""
每日简报生成器 v6 - 全中文 + Word文档
"""

import urllib.request
import urllib.error
import feedparser
import re
import time
import ssl
import smtplib
import os
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.header import Header
from email import encoders
from datetime import datetime, timezone, timedelta
from html import unescape

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TIMEZONE = timezone(timedelta(hours=8))
TIMEOUT = 12

SMTP_CONFIG = {
    "host": "smtp.qq.com",
    "port": 465,
    "user": os.environ.get("SMTP_USER", "your_email@qq.com"),
    "password": os.environ.get("SMTP_PASSWORD", "your_smtp_password"),
    "from_name": "每日简报",
    "to": os.environ.get("SMTP_TO", "your_email@qq.com"),
}

SOURCES = {
    "domestic": [
        {"name": "36氪", "url": "https://36kr.com/feed", "max_items": 5},
        {"name": "IT之家", "url": "https://www.ithome.com/rss/", "max_items": 5},
        {"name": "钛媒体", "url": "https://www.tmtpost.com/rss", "max_items": 5},
        {"name": "少数派", "url": "https://sspai.com/feed", "max_items": 4},
    ],
    "ai": [
        {"name": "量子位", "url": "https://www.qbitai.com/feed/", "max_items": 6},
        {"name": "雷锋网", "url": "https://www.leiphone.com/feed", "max_items": 6},
        {"name": "36氪AI过滤", "url": "https://36kr.com/feed", "max_items": 5},
        {"name": "IT之家AI过滤", "url": "https://www.ithome.com/rss/", "max_items": 5},
        {"name": "钛媒体AI过滤", "url": "https://www.tmtpost.com/rss", "max_items": 4},
    ],
}


def fetch_feed(url, max_items=10):
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"})
        with urllib.request.urlopen(req, timeout=TIMEOUT, context=ctx) as resp:
            raw = resp.read()
    except Exception as e:
        print(f"  [WARN] {url} -> {type(e).__name__}")
        return []
    feed = feedparser.parse(raw)
    items = []
    for entry in feed.entries[:max_items]:
        title = unescape(entry.get("title", "")).strip()
        link = entry.get("link", "")
        summary = unescape(entry.get("summary", entry.get("description", ""))).strip()
        summary = re.sub(r"<[^>]+>", "", summary)
        summary = re.sub(r"\s+", " ", summary).strip()[:200]
        if title:
            items.append({"title": title, "link": link, "summary": summary})
    return items


def deduplicate(items):
    seen = set()
    result = []
    for item in items:
        key = item["title"][:30].lower()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


def is_ai_related(item):
    keywords = [
        "ai", "人工智能", "大模型", "机器学习", "深度学习", "神经网络",
        "gpt", "chatgpt", "openai", "claude", "anthropic", "gemini",
        "机器人", "人形机器人", "自动驾驶", "无人驾驶",
        "芯片", "nvidia", "英伟达", "算法", "模型", "训练", "推理", "算力",
        "数字人", "aigc", "生成式", "多模态", "智能驾驶",
        "llm", "transformer", "diffusion", "agent", "智能体",
    ]
    text = (item["title"] + " " + item["summary"]).lower()
    return any(kw in text.lower() for kw in keywords)


def set_run_font(run, name="微软雅黑", size=Pt(11), color=None, bold=False):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def create_word_doc(domestic, ai_items, filename):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    today = datetime.now(TIMEZONE)
    date_str = today.strftime("%Y年%m月%d日")
    weekday_map = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    weekday_str = weekday_map[today.weekday()]

    title = doc.add_heading(f"每日简报 | {date_str}（{weekday_str}）", level=0)
    for run in title.runs:
        set_run_font(run, size=Pt(22), bold=True)

    doc.add_heading("🔥 国内热点", level=1)
    if domestic:
        for item in domestic[:6]:
            p = doc.add_paragraph()
            run = p.add_run(item['title'])
            set_run_font(run, bold=True, size=Pt(11))
            if item['summary']:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(item['summary'][:120])
                set_run_font(run2, size=Pt(10), color=RGBColor(100, 100, 100))
    else:
        doc.add_paragraph("（暂无数据）")

    doc.add_heading("🤖 AI 领域动态", level=1)
    if ai_items:
        for item in ai_items[:8]:
            p = doc.add_paragraph()
            run = p.add_run(item['title'])
            set_run_font(run, bold=True, size=Pt(11))
            if item['summary']:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(item['summary'][:150])
                set_run_font(run2, size=Pt(10), color=RGBColor(100, 100, 100))
    else:
        doc.add_paragraph("（暂无数据）")

    doc.add_paragraph("")
    p_summary = doc.add_paragraph()
    run_label = p_summary.add_run("📌 一句话总结：")
    set_run_font(run_label, bold=True, size=Pt(11))
    parts = []
    if domestic:
        parts.append(f"国内·{domestic[0]['title'][:40]}")
    if ai_items:
        parts.append(f"AI·{ai_items[0]['title'][:40]}")
    if parts:
        run_content = p_summary.add_run(" | ".join(parts))
        set_run_font(run_content, size=Pt(10))

    doc.add_paragraph("")
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run(f"简报自动生成于 {today.strftime('%H:%M')} | 数据来源：36氪 / IT之家 / 钛媒体 / 少数派 / 量子位 / 雷锋网")
    set_run_font(run_footer, size=Pt(9), color=RGBColor(150, 150, 150))

    doc.save(filename)
    print(f"  ✅ Word文档已保存: {filename}")


def send_email_with_attachment(doc_filename):
    cfg = SMTP_CONFIG
    today = datetime.now(TIMEZONE).strftime("%Y年%m月%d日")
    msg = MIMEMultipart()
    msg["From"] = f"{cfg['from_name']} <{cfg['user']}>"
    msg["To"] = cfg["to"]
    msg["Subject"] = Header(f"📰 每日简报 | {today}", "utf-8")
    body = f"您好，\n\n今日简报已生成，请查看附件。\n\n—— 每日简报自动发送"
    msg.attach(MIMEText(body, "plain", "utf-8"))
    with open(doc_filename, "rb") as f:
        attachment = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
        attachment.set_payload(f.read())
        encoders.encode_base64(attachment)
        attachment.add_header("Content-Disposition", "attachment", filename=Header(os.path.basename(doc_filename), "utf-8").encode())
        msg.attach(attachment)
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP_SSL(cfg["host"], cfg["port"], context=ctx) as server:
            server.login(cfg["user"], cfg["password"])
            server.sendmail(cfg["user"], cfg["to"], msg.as_string())
        print(f"  ✅ 邮件已发送至 {cfg['to']}")
        return True
    except Exception as e:
        print(f"  ❌ 邮件发送失败: {e}")
        return False


def main():
    t0 = time.time()
    print(f"[{datetime.now(TIMEZONE).strftime('%H:%M:%S')}] 开始抓取新闻...\n")
    all_domestic = []
    all_ai = []
    print("[1/2] 国内热点...")
    for src in SOURCES["domestic"]:
        items = fetch_feed(src["url"], src["max_items"])
        print(f"  {src['name']}: {len(items)} 条")
        all_domestic.extend(items)
        time.sleep(0.3)
    print("\n[2/2] AI动态...")
    for src in SOURCES["ai"]:
        items = fetch_feed(src["url"], src["max_items"])
        ai_items = [i for i in items if is_ai_related(i)]
        print(f"  {src['name']}: {len(items)} 条 (AI相关: {len(ai_items)})")
        all_ai.extend(ai_items)
        time.sleep(0.3)
    all_domestic = deduplicate(all_domestic)
    all_ai = deduplicate(all_ai)
    print(f"\n汇总: 国内 {len(all_domestic)} 条, AI {len(all_ai)} 条")
    today = datetime.now(TIMEZONE)
    doc_filename = f"每日简报_{today.strftime('%Y%m%d')}.docx"
    create_word_doc(all_domestic, all_ai, doc_filename)
    elapsed = time.time() - t0
    print(f"\n✅ 简报生成完成 (耗时 {elapsed:.1f}s)")
    print("\n[发送邮件]...")
    send_email_with_attachment(doc_filename)
    print("\n" + "=" * 55)
    print("完成！")


if __name__ == "__main__":
    main()
